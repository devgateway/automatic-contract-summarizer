
import logging
import os
import io
import re
from dataclasses import dataclass
from typing import Optional, List

# --- Permissive (MIT/Apache-2.0) dependencies ---
from bs4 import BeautifulSoup  # MIT
from docx import Document  # python-docx (MIT)
from markdownify import markdownify as _markdownify  # MIT
from pdfminer.high_level import extract_text as _pdfminer_extract_text, extract_text_to_fp as _pdfminer_extract_text_to_fp  # MIT
from pdfminer.layout import LAParams  # MIT
import pdfplumber  # MIT

# Optional helper from your project
try:
    from src.common.html_utils import simplify_html  # project-local
except Exception:
    def simplify_html(html: str) -> str:  # fallback no-op
        return html


# ---------------------------------------------------------------------------
#                      Backwards-compat utility functions
# ---------------------------------------------------------------------------

def remove_commas_from_long_dates(text):
    # Define a regular expression pattern to match the date format "Friday, 19th July, 2024"
    pattern = r"\D*\s(\d{1,2})(th|st|nd|rd)\s\D*,?\s(\d{4})\b"

    def remove_commas(match):
        # Preserve previous debug behavior
        print(match.group(0))
        replaced = match.group(0).replace(",", "")
        print(replaced)
        return replaced

    return re.sub(pattern, remove_commas, text)


def replace_multiple_newlines(text):
    """
    Replace multiple newlines with a single newline. To use less tokens.
    :param text:
    :return:
    """
    temp_text = re.sub(r' +', ' ', text)
    temp_text = re.sub(r'\n+', '\n', temp_text)
    temp_text = re.sub(r'\n+ +', '', temp_text)
    return temp_text


def remove_duplicate_paragraphs(text):
    """
    A side effect of the pdf/docx conversion to text is sometimes entire paragraphs are duplicated. This function
    will remove duplicate paragraphs.
    :param text:
    :return:
    """
    paragraphs = text.split('\\n')
    seen = set()
    unique_paragraphs = []
    for paragraph in paragraphs:
        normalized = paragraph
        if len(normalized) > 15:
            if normalized and normalized not in seen:
                seen.add(normalized)
                unique_paragraphs.append(paragraph)
        else:
            unique_paragraphs.append(normalized)
    return '\\n'.join(unique_paragraphs)


# Keep this for API compatibility; not used by the new DOCX->HTML pipeline.
def ignore_images(element):
    return ""


options = {
    "convert_image": ignore_images
}


# ---------------------------------------------------------------------------
#                                 HTML
# ---------------------------------------------------------------------------

def html_to_text(html_path):
    # print("extract_text_from_html")
    with open(html_path, 'rb') as file:
        html_content = file.read()
        soup = BeautifulSoup(html_content, 'lxml')
        text = soup.get_text()
        # Keep original behavior: strip commas to improve number parsing downstream.
        text = text.replace(",", "")
    return text


def html_to_markdown(html):
    # Convert HTML to Markdown with MIT-licensed markdownify
    return _markdownify(html or "", heading_style="ATX", strip=["style", "script"])


# ---------------------------------------------------------------------------
#                                 DOCX
# ---------------------------------------------------------------------------

def _docx_tables_to_markdown(doc: Document) -> List[str]:
    blocks = []
    for t in doc.tables:
        rows = []
        for row in t.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        if rows:
            header = rows[0]
            aligns = ["---"] * len(header)
            md_table = [f"| {' | '.join(header)} |", f"| {' | '.join(aligns)} |"]
            for r in rows[1:]:
                md_table.append(f"| {' | '.join(r)} |")
            blocks.append("\\n".join(md_table))
    return blocks


def docx_to_text_simple(docx_path, log_source_file=False):
    print(docx_path)
    doc = Document(docx_path)
    parts = []
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            parts.append(p.text.strip())
    text = "\\n".join(parts)
    text = replace_multiple_newlines(text)
    if log_source_file:
        print(text)
    return text


def docx_to_text(docx_path, log_source_file=False, normalize_lowercase=False):
    """
    # Convert .docx content into plain text, including paragraph and tables.
    :param docx_path:
    :param log_source_file:
    :param normalize_lowercase:
    :return: plain text
    """
    print(docx_path)
    doc = Document(docx_path)

    parts = []

    # paragraphs
    for p in doc.paragraphs:
        txt = (p.text or "").strip()
        if txt:
            parts.append(txt)

    # tables -> markdown
    table_blocks = _docx_tables_to_markdown(doc)
    for block in table_blocks:
        parts.append(block)

    result_text = "\\n".join(parts)
    result_text = replace_multiple_newlines(result_text)
    result_text = remove_duplicate_paragraphs(result_text)
    if normalize_lowercase:
        result_text = result_text.lower()

    if log_source_file:
        print(result_text)
    return result_text


def docx_to_html(file_path, log_source_file=False):
    # Replaces mammoth (BSD-2) with MIT/Apache-only stack using BeautifulSoup post-processing.
    # Prefer PyDocX (Apache Software License) if available.
    print(file_path)
    try:
        from pydocx import PyDocX  # Apache Software License
        html = PyDocX.to_html(file_path)
    except Exception:
        # Fallback: very simple HTML constructed from python-docx contents (headings not preserved)
        doc = Document(file_path)
        parts = ["<div>"]
        for p in doc.paragraphs:
            txt = (p.text or "").strip()
            if txt:
                parts.append(f"<p>{txt}</p>")
        # tables
        for t in doc.tables:
            parts.append("<table>")
            for row in t.rows:
                parts.append("<tr>")
                for cell in row.cells:
                    parts.append(f"<td>{(cell.text or '').strip()}</td>")
                parts.append("</tr>")
            parts.append("</table>")
        parts.append("</div>")
        html = "".join(parts)

    html = simplify_html(html)
    if log_source_file:
        print(html)
    return html


# ---------------------------------------------------------------------------
#                                  PDF
# ---------------------------------------------------------------------------

@dataclass
class _PDFOpts:
    line_margin: float = 0.5
    word_margin: float = 0.1
    char_margin: float = 1.0
    detect_vertical: bool = False
    include_tables: bool = True  # append tables as markdown blocks


def _extract_pdf_tables_as_markdown(pdf_path: str) -> List[str]:
    blocks: List[str] = []
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for pi, page in enumerate(pdf.pages, start=1):
                page_tables = []
                # Try both line and text strategies for better recall
                for setting in (
                    {"vertical_strategy": "lines", "horizontal_strategy": "lines"},
                    {"vertical_strategy": "text", "horizontal_strategy": "text"},
                ):
                    try:
                        page_tables.extend(page.extract_tables(table_settings=setting) or [])
                    except Exception:
                        continue

                # Deduplicate
                seen = set()
                uniq = []
                for tbl in page_tables:
                    key = tuple(tuple((c or "") for c in row) for row in tbl)
                    if key not in seen:
                        seen.add(key)
                        uniq.append(tbl)

                if not uniq:
                    continue

                md_blocks = []
                for t in uniq:
                    header = t[0]
                    aligns = ["---"] * len(header)
                    md_table = [f"| {' | '.join((c or '') for c in header)} |",
                                f"| {' | '.join(aligns)} |"]
                    for r in t[1:]:
                        md_table.append(f"| {' | '.join((c or '') for c in r)} |")
                    md_blocks.append("\\n".join(md_table))

                blocks.append(f"\\n\\n## [PDF Tables - page {pi}]\\n\\n" + "\\n\\n".join(md_blocks))
    except Exception as e:
        # Keep resilient behavior
        logging.getLogger(__name__).warning("Table extraction failed: %s", e)
    return blocks


def pdf_to_text(pdf_path, log_source_file=False, normalize_lowercase=False):
    """
    Convert .pdf file to plain text, including paragraph and tables.
    :param pdf_path:
    :param log_source_file:
    :param normalize_lowercase:
    :return:
    """
    print(pdf_path)

    # 1) base text via pdfminer.six
    laparams = LAParams(line_margin=0.5, word_margin=0.1, char_margin=1.0, detect_vertical=False, all_texts=True)
    base_text = _pdfminer_extract_text(pdf_path, laparams=laparams) or ""

    # 2) tables via pdfplumber -> markdown blocks
    table_blocks = _extract_pdf_tables_as_markdown(pdf_path)
    result_text = base_text + ("".join(table_blocks) if table_blocks else "")

    result_text = replace_multiple_newlines(result_text)
    result_text = remove_duplicate_paragraphs(result_text)
    if normalize_lowercase:
        result_text = result_text.lower()

    if log_source_file:
        print(result_text)
    return result_text


def pdf_to_text_with_intermediate_docx(pdf_path, log_source_file=False):
    print(pdf_path)
    # Backward-compatible shim: previously used pdf2docx + mammoth. Now use direct, permissive-only stack.
    text = pdf_to_text(pdf_path, log_source_file=log_source_file, normalize_lowercase=False)
    if log_source_file:
        print(text)
    return text


def pdf_to_html(pdf_path, log_source_file=False):
    print(pdf_path)
    output = io.StringIO()
    laparams = LAParams()
    with open(pdf_path, "rb") as f:
        _pdfminer_extract_text_to_fp(f, output, laparams=laparams, output_type="html", codec=None)
    html = output.getvalue()
    html = simplify_html(html)
    if log_source_file:
        print(html)
    return html


# ---------------------------------------------------------------------------
#                            Misc project helpers
# ---------------------------------------------------------------------------

def extract_fields_from_text(json_data):
    # print("extract_fields_from_text")
    # Extract the contract title, ocid and budget amount from the text
    fields = (json_data.get("releases")[0].get("tender").get("title")
              + '\\n' + json_data.get("releases")[0].get("ocid") + '\\n'
              + str(json_data.get("releases")[0].get("planning").get("budget").get("amount").get("amount")))
    return fields
